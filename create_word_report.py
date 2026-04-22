import os
import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('DocuChain: Nexus Protocol V2.0\nProject Report', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Enterprise-Grade Document Verification, Anti-Malware Sandboxing, and Blockchain-Anchored Storage System.').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph('')

doc.add_heading('1. Abstract / Introduction', level=1)
doc.add_paragraph("DocuChain is a comprehensive Zero-Trust architecture designed to secure organizational document workflows. It intercepts untrusted files at the browser level, sanitizes them via multi-engine malware scanning, isolates them in a secure sandbox, and anchors verified assets into an immutable Hyperledger Fabric blockchain ledger backed by private MinIO object storage.")

doc.add_heading('2. Core Architecture', level=1)
doc.add_paragraph("DocuChain transitions from a public Web3 stack to a Permissioned Enterprise Architecture.")
doc.add_paragraph("Frontend Client: React.js, Tailwind CSS, Vite. Features a highly interactive, futuristic \"Glassmorphism\" UI with real-time state management.", style='List Bullet')
doc.add_paragraph("Security Extension: Manifest V3 Chrome Extension. Operates as a background service worker intercepting browser-level file I/O operations dynamically.", style='List Bullet')
doc.add_paragraph("Backend Server: FastAPI (Python). High-performance asynchronous API handling authentication, document routing, and MinIO/Blockchain orchestration.", style='List Bullet')
doc.add_paragraph("Object Storage: MinIO. Self-hosted, S3-compatible enterprise storage. Assets are segregated into Private Hub buckets and a Global Audit bucket.", style='List Bullet')
doc.add_paragraph("Blockchain Ledger: Hyperledger Fabric. Permissioned ledger ensuring cryptographic immutability of document hashes (SHA-256) and transaction timestamps.", style='List Bullet')

doc.add_heading('3. Detailed Component Breakdown', level=1)
doc.add_heading('3.1 The Browser Security Extension', level=2)
doc.add_paragraph("The extension acts as the first line of defense (The \"Arise\" Protocol).")
doc.add_paragraph("Pre-Download Interception: Hooks into chrome.downloads to aggressively pause downloads before they hit the local disk.", style='List Bullet')
doc.add_paragraph("Multi-Engine Scanning: Relays file buffers or download URLs to the backend, which queries VirusTotal and Internxt APIs for zero-day malware signatures.", style='List Bullet')
doc.add_paragraph("Scanner UI: Displays a dynamic \"Radar\" interface providing real-time feedback on the scanning process.", style='List Bullet')

doc.add_heading('3.2 DocuVault (Private Hubs)', level=2)
doc.add_paragraph("A multi-tenant architecture allowing organizations to create isolated workspaces (Hubs).")
doc.add_paragraph("Role-Based Access: Users can generate secure invite codes to add members to their private network hub.", style='List Bullet')
doc.add_paragraph("Asset Management: Users can view their privately anchored assets, generate Temporary Access Tokens (24h expiry), or generate Offline Verification QR Codes.", style='List Bullet')
doc.add_paragraph("Asset Destruction: Complete binary wipe from MinIO storage and SQLite metadata erasure for complete data sovereignty.", style='List Bullet')

doc.add_heading('3.3 Global Audit Trail', level=2)
doc.add_paragraph("A transparent ledger for cross-organization verification.")
doc.add_paragraph("Asset Promotion: Users can push private assets to the Global Vault, moving the binary across MinIO buckets.", style='List Bullet')
doc.add_paragraph("Immutability Score: Displays the cryptographic integrity score comparing the live MinIO binary hash against the Hyperledger Fabric anchor.", style='List Bullet')
doc.add_paragraph("Owner Sovereignty: Only the original uploader retains the cryptographic authority to delete the asset from the Global Audit ledger.", style='List Bullet')

doc.add_heading('4. Security & Zero-Trust Protocol', level=1)
doc.add_paragraph("Blob Bypass Architecture: Prevents the extension from recursively intercepting its own sandbox blob URIs.", style='List Bullet')
doc.add_paragraph("JWT Session Management: Stateless authentication using Bearer tokens ensuring strict network isolation.", style='List Bullet')
doc.add_paragraph("Cryptographic Verification: API calculates the hash of the currently stored file and compares it against the initial hash on the Fabric ledger.", style='List Bullet')

doc.add_heading('5. System Workflows', level=1)
doc.add_paragraph("Note: Refer to the README.md for the Mermaid Flowchart diagrams covering Security Interception, Asset Anchoring, and Asset Promotion.")

doc.save('DocuChain_Project_Report.docx')
print("Successfully generated DocuChain_Project_Report.docx")
